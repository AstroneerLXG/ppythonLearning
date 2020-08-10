from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 用于段落对齐操作
from docx.shared import Pt  # 磅数
from docx.oxml.ns import qn  # 中文格式
# 以上是docx库中所需部分
import time

priceToday = input('请输入今日价格：')
companyList = ['西北工业大学', '西安交通大学', '西安电子科技大学', '西北大学']
# dateToday = time.strftime('%Y/%m/%d', time.localtime())  # 对中文支持较差，只能用英文字符串
dateToday = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日')  # 更好地支持中文

for i in companyList:
    newDocument = Document()
    newDocument.styles['Normal'].font.name = u'宋体'  # 设置英文和数字的默认字体（由于对中文支持较差，这里对中文不起作用）
    newDocument.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文默认字体

    para1 = newDocument.add_paragraph()  # 第一段（标题）
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐，默认为左对齐
    para1.space_before = Pt(5)  # 段前段后间距
    para1.space_after = Pt(5)
    run1 = para1.add_run('关于下达%s产品价格的通知' % dateToday)  # 填写内容
    run1.font.name = '等线'  # 设置英文字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')  # 设置中文字体
    run1.font.size = Pt(21)
    run1.font.bold = True

    para2 = newDocument.add_paragraph()  # 第二段（客户名称）
    para2.space_before = Pt(5)
    run2 = para2.add_run(i + '：')
    run2.font.size = Pt(16)
    run2.font.bold = True

    para3 = newDocument.add_paragraph()  # 第三段（正文）
    run3 = para3.add_run('    根据公司安排，今日x86芯片价格已拟定，为%s元，特此通知，感谢您的选择！' % priceToday)
    run3.font.size = Pt(16)

    para4 = newDocument.add_paragraph()  # 第四段（落款）
    para4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
    run4 = para4.add_run('联系人：小李    电话：12345678901')
    run4.font.size = Pt(14)
    run2.font.bold = True

    newDocument.save('C:/Users/aby/Desktop/%s-价格通知.docx' % i)
